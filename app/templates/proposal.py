"""Render a structured Proposal into a downloadable .docx file."""

from __future__ import annotations

import io

from docx import Document

from app.schemas.rfp import Proposal


def render_proposal(proposal: Proposal) -> bytes:
    """Render a Proposal (sections + extraction metadata) to .docx bytes."""
    extraction = proposal.extraction
    document = Document()

    document.add_heading("Project Proposal", level=0)

    if extraction.project_title:
        document.add_paragraph(f"In response to: {extraction.project_title}")
    if extraction.company_name:
        document.add_paragraph(f"Prepared for: {extraction.company_name}")
    if extraction.submission_deadline:
        document.add_paragraph(f"Submission deadline: {extraction.submission_deadline}")
    if extraction.budget_range:
        document.add_paragraph(f"Budget range: {extraction.budget_range}")

    if proposal.sections:
        document.add_page_break()

    for section in proposal.sections:
        document.add_heading(section.title, level=1)
        for paragraph_text in section.content.split("\n\n"):
            paragraph_text = paragraph_text.strip()
            if paragraph_text:
                document.add_paragraph(paragraph_text)

    if extraction.flags:
        document.add_page_break()
        document.add_heading("Internal Review Notes", level=1)
        document.add_paragraph(
            "The following notes were generated during automated extraction "
            "and should be reviewed before submission."
        )
        for flag in extraction.flags:
            document.add_paragraph(flag, style="List Bullet")

    buffer = io.BytesIO()
    document.save(buffer)
    return buffer.getvalue()
